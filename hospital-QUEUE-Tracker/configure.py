from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word Document
doc = Document()

# --- Function to add a heading and content ---
def add_section(title, content):
    doc.add_heading(title, level=1)
    for paragraph in content:
        p = doc.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(12)
    doc.add_page_break()

# --- Cover Page ---
cover = doc.add_paragraph()
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cover.add_run("\n\nHospital Queue Tracker\n").bold = True
cover.add_run("Web & Mobile Hybrid Application\n")
cover.add_run("Submitted By: Gopi Nath P\n")
cover.add_run("Course/Department: [Your Department Name]\n")
cover.add_run("Institution: [Your College/Institution Name]\n")
cover.add_run("Date: [Submission Date]\n")
cover.add_run("[Insert hospital/healthcare logo]\n")
doc.add_page_break()

# --- Certificate / Acknowledgement ---
cert_content = [
    "This is to certify that the project titled 'Hospital Queue Tracker' is the original work of Gopi Nath P, "
    "submitted in partial fulfillment of the requirements for the award of [Degree Name] at [Institution Name].",
    "Acknowledgement: I would like to express my sincere gratitude to my project guide [Mentor Name] for providing guidance "
    "and support throughout this project. I also acknowledge my friends and family for their encouragement and motivation."
]
add_section("Certificate / Acknowledgement", cert_content)

# --- Table of Contents ---
toc = [
    "1. Introduction", "2. Objectives", "3. Problem Definition", "4. Scope of the Project",
    "5. Literature Survey / Related Works", "6. System Design & Architecture",
    "7. Requirement Analysis", "8. Database Design", "9. Implementation",
    "10. Testing", "11. Results & Discussion", "12. Future Enhancements",
    "13. Conclusion", "14. References", "15. Annexure"
]
add_section("Table of Contents", toc)

# --- Full Content Sections ---
sections = {
    "Introduction": [
        "Hospitals often face overcrowding in outpatient departments.",
        "Manual queue systems cause long waiting times and patient frustration.",
        "Digital queue trackers improve workflow efficiency and reduce waiting time.",
        "Key benefits include real-time patient tracking, automated token management, reduced errors in patient handling, improved patient satisfaction.",
        "Real-world example: City hospitals using token-based mobile applications have shown measurable improvement in waiting time and operational efficiency.",
        "[Placeholder: Hospital Queue Workflow Diagram]"
    ],
    "Objectives": [
        "Reduce patient waiting time using digital queue management.",
        "Provide real-time updates on queue status for doctors, patients, and admin.",
        "Implement separate modules for Admin, Doctor, and Patient.",
        "Enable automatic token generation and status tracking.",
        "Maintain a record of patient flow for analysis and reporting.",
        "Ensure scalability for different hospital sizes.",
        "Provide future scope for mobile integration and AI predictions."
    ],
    "Problem Definition": [
        "Overcrowding in hospitals leads to patient dissatisfaction.",
        "Manual queue systems result in mismanagement of tokens, incorrect patient tracking, long waiting times.",
        "Doctors face difficulty in managing patient flow efficiently.",
        "Administrators have no real-time dashboard to track the queue.",
        "Existing digital systems may lack proper analytics or mobile integration."
    ],
    "Scope of the Project": [
        "Applicable to small, medium, and large hospitals.",
        "Can integrate with mobile notifications for patients.",
        "Can scale for multiple departments within the hospital.",
        "Future integration possibilities: AI-based wait-time predictions, cloud database storage, integration with hospital management software.",
        "Provides dashboards for monitoring, reporting, and analytics."
    ],
    "Literature Survey / Related Works": [
        "Current queue management systems include manual token systems, digital queue apps, and hospital management software with queue modules.",
        "Manual systems: Low cost, simple, but inefficient and error-prone.",
        "Digital systems: Real-time, accurate, provides analytics, but requires setup.",
        "Comparative study shows digital systems significantly reduce waiting time.",
        "Gap analysis: Many systems lack real-time dashboards, modular design, and patient-friendly interface.",
        "[Placeholder: Comparative table of existing systems]"
    ],
    "System Design & Architecture": [
        "Modules: Admin (manage doctors, departments, monitor queue), Doctor (view token list, call next patient, update status), Patient (generate token, check queue).",
        "Workflow: Patient generates token → Doctor views queue and calls next patient → Admin monitors overall queue.",
        "Frontend communicates with Flask backend; Backend interacts with MySQL database.",
        "Security: Role-based access, password encryption, session management.",
        "[Placeholders: System Architecture Diagram, Data Flow Diagram, UML Diagrams (Use Case, Class, Sequence)]"
    ],
    "Requirement Analysis": [
        "Functional Requirements:",
        "Admin: Login securely, Add/Edit/Remove doctors, Monitor queues and update status.",
        "Doctor: Login securely, View tokens and call next patient, Update patient treatment status.",
        "Patient: Generate token, View queue status and estimated wait time.",
        "Non-Functional Requirements: User-friendly interface, Performance: handle 100+ users, Security: encryption, Scalability.",
        "Software Requirements: Python, Flask, MySQL, VS Code.",
        "Hardware Requirements: PC with 8GB RAM, Mobile device for testing."
    ],
    "Database Design": [
        "ER Diagram: Relationships between Patients, Doctors, and Departments.",
        "Tables: Patients (id, name, phone, token, status), Doctors (id, name, department), Departments (id, name).",
        "[Placeholder: ER Diagram]"
    ],
    "Implementation": [
        "Frontend: HTML, CSS, JavaScript.",
        "Backend: Python Flask.",
        "Database: MySQL.",
        "Module Implementation: Admin Dashboard, Doctor Dashboard, Patient interface.",
        "[Placeholder: Screenshots of all modules]"
    ],
    "Testing": [
        "Unit Testing: Test individual modules separately.",
        "Integration Testing: Test end-to-end workflow.",
        "Functional Testing: Verify Admin, Doctor, Patient functionalities.",
        "Sample Test Case Table: Test Case | Input | Expected Output | Actual Output | Status",
        "[Placeholder: Screenshots of test results]"
    ],
    "Results & Discussion": [
        "Screenshots of final system dashboards.",
        "Comparison with manual queue system.",
        "Benefits: Reduced waiting time, improved workflow, real-time monitoring.",
        "Observations: System is user-friendly, scalable, and effective."
    ],
    "Future Enhancements": [
        "Mobile app integration.",
        "SMS and Email notifications.",
        "AI-based wait-time prediction.",
        "Cloud database integration.",
        "Integration with hospital management systems."
    ],
    "Conclusion": [
        "Objectives achieved: Reduced waiting time and improved patient experience.",
        "Scalable and user-friendly system.",
        "Ready for deployment in hospitals.",
        "Provides foundation for further enhancements."
    ],
    "References": [
        "Books, research papers, and websites.",
        "APA / IEEE citation style.",
        "[Placeholder: List of references]"
    ],
    "Annexure": [
        "Code snippets (if needed).",
        "Screenshots of login, admin dashboard, doctor dashboard, patient token pages.",
        "Database tables with sample data.",
        "UML and ER diagrams placeholders."
    ]
}

# Add all sections to document
for title, content in sections.items():
    add_section(title, content)

# --- Save Word Document ---
file_path = "Hospital_Queue_Tracker_Expanded_Report.docx"
doc.save(file_path)

print(f"✅ Full expanded Word report generated: {file_path}")
        